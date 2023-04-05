import os
# from flask import Flask, request, make_response, jsonify
from docx import Document
import json
from docx.enum.text import WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from htmldocx import HtmlToDocx
from flask_cors import CORS
import time
import sys


# app = Flask(__name__)
# CORS(app)

def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    # paragraph.style.font.size = Pt(5)


# @app.route('/download_cv', methods=['POST'])
def download_docx(data):
    # create a new document

    # from data import data  # hardcoded data
    # data = json.loads(request.data)  # API data

    doc = Document()
    new_parser = HtmlToDocx()

    # font style and size
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(20)

    name_size = 14
    head_size = 12
    low_head_size = 11
    content_size = 10.5

    # Add full name
    p = doc.add_paragraph()
    p.add_run(f"{data.get('first_name', ' ')} {data.get('last_name', ' ')}").bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.style.font.size = Pt(name_size)

    # Add personal_info
    p = doc.add_paragraph()
    p.add_run(f"{data.get('phone', ' ')} | {data.get('location', ' ')} | {data.get('email', ' ')}").bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.style.font.size = Pt(head_size)

    # Add Professional Summary
    summary_p = doc.add_paragraph('Professional Summary')
    summary_p.bold = True
    summary_p.style.font.size = Pt(head_size)
    insertHR(summary_p)
    summary_dp = doc.add_paragraph(data.get('description', ' '))
    summary_dp.style.font.size = Pt(content_size)
    doc.add_paragraph()

    # Add Professional Skills
    p = doc.add_paragraph('Professional Skills')
    p.bold = True
    p.style.font.size = Pt(head_size)
    insertHR(p)
    p = doc.add_paragraph()
    p.add_run(f'{_  }  |' for _ in data.get('skills', []))
    p.style.font.size = Pt(content_size)
    doc.add_paragraph()

    # Add Work History
    p = doc.add_paragraph('Work History')
    p.bold = True
    p.style.font.size = Pt(head_size)
    insertHR(p)
    for company in data.get('work_history'):
        t_period = company.get('start_date', ' ') + company.get('end_date', ' ')
        time_period = doc.add_paragraph(t_period)
        time_period.add_run('\t')

        time_period.add_run(company.get('job_title', ' ') + ', ')
        time_period.add_run(company.get('company_name', ' '))
        p.add_run().add_break(WD_BREAK.LINE)
        resp = doc.add_paragraph()
        u_b1 = resp.add_run('Responsibilities')
        u_b1.bold = True
        u_b1.underline = True
        new_parser.add_html_to_document(company.get('responsibilities'), doc)
        resp.add_run().add_break(WD_BREAK.LINE)

        ach = doc.add_paragraph()
        u_b2 = ach.add_run('Key Achievements (')
        u_b2.bold = True
        u_b2.underline = True
        u_b3 = ach.add_run('SPECULATIVE')
        u_b3.bold = True
        u_b3.underline = True
        u_b3.font.color.rgb = RGBColor(255, 0, 0)
        u_b4 = ach.add_run(')')
        u_b4.bold = True
        u_b4.underline = True
        ach.add_run().add_break(WD_BREAK.LINE)

        new_parser.add_html_to_document(company.get('achievements'), doc)

        doc.add_paragraph()

    # Add Education and Qualifications
    edu = doc.add_paragraph('Education and Qualifications')
    edu.bold = True
    edu.style.font.size = Pt(head_size)
    insertHR(edu)
    for company in data.get('education'):
        time_period2 = doc.add_paragraph(company.get('degree_type', ' ') + '  ')
        time_period2.add_run('\t')

        time_period2.add_run(company.get('field_of_study', ' ') + ', ')
        time_period2.add_run(company.get('institution_name', ' ') + ', ')
        t_period = company.get('start_date', ' ') + company.get('end_date', ' ')
        time_period2.add_run(t_period)
        time_period2.add_run().add_break(WD_BREAK.LINE)
        new_parser.add_html_to_document(company.get('description'), doc)
        doc.add_paragraph()

    # Add Training and Certification
    edu = doc.add_paragraph('Training and Certification')
    edu.bold = True
    edu.style.font.size = Pt(head_size)
    insertHR(edu)
    for cert in data.get('certificates'):
        time_period2 = doc.add_paragraph(cert.get('degree_type', ' ') + ', ')
        time_period2.add_run('\t')

        time_period2.add_run(cert.get('institution_name', ' ') + ', ')
        time_period2.add_run(cert.get('start_date', ' '))
        time_period2.add_run().add_break(WD_BREAK.LINE)
        new_parser.add_html_to_document(cert.get('description'), doc)
        doc.add_paragraph()

    # Add Awards and Recognitions
    edu = doc.add_paragraph('Awards and Recognitions')
    edu.bold = True
    edu.style.font.size = Pt(head_size)
    insertHR(edu)
    for cert in data.get('awards'):
        time_period2 = doc.add_paragraph(cert.get('volunteering_title', ' ') + ', ')
        time_period2.add_run(cert.get('organisation_name', ' ') + ', ')
        t_period = cert.get('start_date', ' ') + cert.get('end_date', ' ')
        time_period2.add_run(t_period)
        doc.add_paragraph()

    # Add Publications and Presentations
    edu = doc.add_paragraph('Publications and Presentations')
    edu.bold = True
    edu.style.font.size = Pt(head_size)
    insertHR(edu)
    for cert in data.get('publications'):
        time_period2 = doc.add_paragraph(cert.get('title', ' ') + ', ')
        time_period2.add_run(cert.get('publisher', ' ') + ', ')
        t_period = cert.get('start_date', ' ') + cert.get('end_date', ' ')
        time_period2.add_run(t_period)
        doc.add_paragraph()

    # Add Patents
    edu = doc.add_paragraph('Patents')
    edu.bold = True

    edu.style.font.size = Pt(head_size)
    insertHR(edu)
    for cert in data.get('patents'):
        time_period2 = doc.add_paragraph(cert.get('patent_name', ' ') + ', ')
        time_period2.add_run(cert.get('patent_number', ' ') + ', ')
        t_period = cert.get('start_date', ' ') + cert.get('end_date', ' ')
        time_period2.add_run(t_period)
        time_period2.add_run().add_break(WD_BREAK.LINE)
        time_period2.add_run(cert.get('patent_url', ' '))
        time_period2.add_run()
        new_parser.add_html_to_document(cert.get('description'), doc)
        doc.add_paragraph()

    file_path = r'resume.docx'
    if os.path.isfile(file_path):
        os.remove(file_path)
    doc.save(file_path)
    # response = {'link': 'file://'+os.path.abspath(file_path)}
    # return jsonify(response)
    return 'file://'+os.path.abspath(file_path)


if __name__ == '__main__':
    output = download_docx(json.loads(sys.argv[1]))
    print(output)
